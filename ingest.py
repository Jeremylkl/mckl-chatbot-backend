import os
from pathlib import Path
from typing import List

from dotenv import load_dotenv

import pandas as pd
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
import docx

from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document

load_dotenv()

DATA_DIR = Path("data")
CHROMA_DIR = Path("chroma_db")

CHUNK_SIZE = int(os.getenv("CHUNK_SIZE", 1000))
CHUNK_OVERLAP = int(os.getenv("CHUNK_OVERLAP", 150))


def load_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def load_pdf(path: Path) -> str:
    txt_pages = []
    reader = PdfReader(str(path))
    for page in reader.pages:
        txt_pages.append(page.extract_text() or "")
    text = "\n".join(txt_pages)
    if not text.strip():
        text = load_pdf_with_ocr(str(path))
    return text


def load_pdf_with_ocr(file_path: str) -> str:
    try:
        pages = convert_from_path(file_path)
        text = ""
        for page in pages:
            text += pytesseract.image_to_string(page)
        return text
    except Exception as e:
        print(f"⚠️ OCR failed for {file_path}: {e}")
        return ""


def load_docx(path: Path) -> str:
    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n".join(paragraphs)


def load_documents_from_folder(folder: Path) -> List[Document]:
    docs: List[Document] = []
    for path in sorted(folder.glob("*")):
        if path.is_dir():
            continue

        suffix = path.suffix.lower()
        try:
            if suffix == ".txt":
                text = load_txt(path)
            elif suffix == ".pdf":
                text = load_pdf(path)
            elif suffix in [".docx", ".doc"]:
                text = load_docx(path)
            elif suffix == ".csv":
                df = pd.read_csv(path)
                text = df.to_string(index=False)
            else:
                print(f"Skipping unsupported file type: {path.name}")
                continue

            if not text.strip():
                print(f"Warning: empty content for {path.name}")
                continue

            meta = {"source": str(path)}
            docs.append(Document(page_content=text, metadata=meta))
            print(f"Loaded {path.name}")
        except Exception as e:
            print(f"Error loading {path.name}: {e}")

    return docs


def chunk_documents(docs: List[Document]):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
    )
    return splitter.split_documents(docs)


def create_and_save_vectorstore(docs: List[Document], db_dir: Path):
    embeddings = OpenAIEmbeddings()
    vectordb = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory=str(db_dir),
    )
    vectordb.persist()
    print(f"✅ Saved Chroma DB to {db_dir}")


def main():
    if not DATA_DIR.exists():
        print("❗ Create a 'data/' folder and drop pdf/docx/txt/csv notes there.")
        return

    print("📥 Loading docs...")
    docs = load_documents_from_folder(DATA_DIR)
    if not docs:
        print("❗ No documents found in data/. Put files there and try again.")
        return

    print("✂️ Chunking docs...")
    chunks = chunk_documents(docs)
    print(f"Total chunks: {len(chunks)}")

    print("🧠 Creating vector store...")
    create_and_save_vectorstore(chunks, CHROMA_DIR)
    print("✅ Ingestion complete.")


if __name__ == "__main__":
    main()
